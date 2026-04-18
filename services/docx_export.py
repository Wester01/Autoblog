from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor
from pathlib import Path


def docx_export(post: dict, path: Path):
    document = Document()

    #H1
    titulo = document.add_heading(post["h1"], level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #Secciones
    for seccion in post["secciones"]:
        document.add_heading(seccion["h2"], level=2)

        parrafo = document.add_paragraph()

        partes = seccion["contenido"].split("**")

        for i, parte in enumerate(partes):
            run = parrafo.add_run(parte)

            if i % 2 != 0:
                run.bold = True

    document.save(str(path))